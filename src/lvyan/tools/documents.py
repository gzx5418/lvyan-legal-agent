"""文档工具：文档文本提取与合同条款风险分析。

本模块是 SubTask 15.3 的实现，提供两个标准工具：

  - ``extract_document(filepath)``：从 .docx / .md / .txt 文件提取文本。
  - ``analyze_contract_clause(text, clause_type=...)``：基于关键词规则检测
    常见合同风险条款（桩实现）。

TODO: ``analyze_contract_clause`` 后续接入 LLM 分析。
"""

from __future__ import annotations

from datetime import datetime
from pathlib import Path
from typing import Literal

from pydantic import Field

from lvyan.tools.base import ToolResult
from lvyan.validators.prompt_injection import (
    InjectionDetectionResult,
    detect_prompt_injection,
)

# ---------------------------------------------------------------------------
# 常量
# ---------------------------------------------------------------------------
_TEXT_EXCERPT_LIMIT = 5000
"""DocumentExtractResult.text 截断长度（字符数）。"""

DocType = Literal["docx", "md", "txt", "unknown"]


# 风险条款关键词映射：风险点 -> 关键词列表
_RISK_CLAUSE_PATTERNS: dict[str, list[str]] = {
    "定金不退": ["定金不退", "定金不予退还", "定金概不退还", "定金不予返还"],
    "最终解释权": ["最终解释权", "保留最终解释权", "享有最终解释权"],
    "单方变更": [
        "单方变更",
        "单方修改",
        "有权随时变更",
        "有权随时修改",
        "无需通知变更",
    ],
    "加重对方责任": [
        "违约金按",
        "每日按",
        "承担一切责任",
        "承担全部责任",
        "无限连带责任",
    ],
    "限制对方权利": [
        "不得主张",
        "不得要求",
        "不得解除",
        "不得索赔",
        "视为放弃",
        "视为自动放弃",
    ],
    "管辖不利": [
        "由甲方所在地",
        "由乙方所在地",
        "必须在我方所在地",
        "仲裁委员会（如有）",
    ],
    "自动续约": ["自动续约", "自动续期", "默认续约", "无需通知续约"],
    "没收财产": ["没收", "扣除全部", "不予退还任何费用"],
    "格式条款排除": ["已充分理解", "已仔细阅读", "自愿接受全部条款", "不再主张任何权利"],
}
"""常见合同风险条款关键词映射。"""


# ---------------------------------------------------------------------------
# 返回模型
# ---------------------------------------------------------------------------
class DocumentExtractResult(ToolResult):
    """文档文本提取结果。

    新增 ``injection_detection`` 字段（Task 18.1）：对提取出的全文做提示注入
    检测，检测到时 ``injection_detection.detected=True`` 并附 ``warning``，
    但**不修改** ``text`` 字段（仅标记，由调用方决定是否拒绝）。
    """

    filepath: str
    doc_type: DocType = "unknown"
    text: str = ""  # excerpt 前 5000 字
    full_text_length: int = 0
    extracted_at: datetime = Field(default_factory=datetime.now)
    injection_detection: InjectionDetectionResult | None = None
    # success / error 由基类提供


class ContractAnalysisResult(ToolResult):
    """合同条款风险分析结果。"""

    clause_type: str | None = None
    risk_level: Literal["low", "medium", "high"] = "low"
    risk_factors: list[str] = Field(default_factory=list)
    suggestions: list[str] = Field(default_factory=list)
    analyzed_text_length: int = 0


# ---------------------------------------------------------------------------
# 公开工具
# ---------------------------------------------------------------------------
def extract_document(filepath: str) -> DocumentExtractResult:
    """从 .docx / .md / .txt 文件提取文本。

    Args:
        filepath: 文件路径（.docx 用 python-docx；.md / .txt 直接读取）。

    Returns:
        DocumentExtractResult：text 字段为正文前 5000 字，full_text_length 为
        完整文本长度。失败时 success=False 且 error 含原因。
    """
    if not filepath or not filepath.strip():
        return DocumentExtractResult(
            tool_name="extract_document",
            success=False,
            error="filepath 不能为空",
            filepath=filepath or "",
        )

    path = Path(filepath)
    if not path.is_file():
        return DocumentExtractResult(
            tool_name="extract_document",
            success=False,
            error=f"文件不存在：{filepath}",
            filepath=filepath,
        )

    suffix = path.suffix.lower()
    extracted_at = datetime.now()

    try:
        if suffix == ".docx":
            text = _extract_docx(path)
            doc_type: DocType = "docx"
        elif suffix in (".md", ".markdown"):
            text = _read_text_safely(path)
            doc_type = "md"
        elif suffix == ".txt":
            text = _read_text_safely(path)
            doc_type = "txt"
        else:
            # 兜底：尝试按文本读取
            text = _read_text_safely(path)
            doc_type = "unknown"
    except Exception as exc:  # noqa: BLE001
        return DocumentExtractResult(
            tool_name="extract_document",
            success=False,
            error=f"提取失败：{exc}",
            filepath=filepath,
            extracted_at=extracted_at,
        )

    excerpt = text[:_TEXT_EXCERPT_LIMIT]

    # 提示注入检测（Task 18.1）：对提取出的正文做标记型检测，检测到注入时
    # 仅在 ``injection_detection`` 中标记 warning，不修改 ``text`` 字段，
    # 由调用方决定是否拒绝。检测异常不中断提取流程。
    try:
        injection_detection = detect_prompt_injection(excerpt)
    except Exception:  # noqa: BLE001 检测异常不中断提取
        injection_detection = None

    return DocumentExtractResult(
        tool_name="extract_document",
        success=True,
        filepath=filepath,
        doc_type=doc_type,
        text=excerpt,
        full_text_length=len(text),
        extracted_at=extracted_at,
        injection_detection=injection_detection,
    )


def analyze_contract_clause(
    text: str,
    clause_type: str | None = None,
) -> ContractAnalysisResult:
    """分析合同条款风险（桩实现：基于关键词规则）。

    Args:
        text: 合同条款文本。
        clause_type: 条款类型提示（如「违约责任」「管辖」），可选。

    Returns:
        ContractAnalysisResult：含 risk_level / risk_factors / suggestions。
    """
    if not text or not text.strip():
        return ContractAnalysisResult(
            tool_name="analyze_contract_clause",
            success=False,
            error="待分析文本不能为空",
            clause_type=clause_type,
            analyzed_text_length=0,
        )

    detected: list[str] = []
    for risk_name, keywords in _RISK_CLAUSE_PATTERNS.items():
        for kw in keywords:
            if kw in text:
                detected.append(risk_name)
                break

    risk_level: Literal["low", "medium", "high"]
    if len(detected) >= 3:
        risk_level = "high"
    elif len(detected) >= 1:
        risk_level = "medium"
    else:
        risk_level = "low"

    suggestions = _build_suggestions(detected, clause_type)

    # TODO: 后续接入 LLM 分析
    return ContractAnalysisResult(
        tool_name="analyze_contract_clause",
        success=True,
        clause_type=clause_type,
        risk_level=risk_level,
        risk_factors=detected,
        suggestions=suggestions,
        analyzed_text_length=len(text),
    )


# ---------------------------------------------------------------------------
# 内部辅助
# ---------------------------------------------------------------------------
def _extract_docx(path: Path) -> str:
    """用 python-docx 提取 .docx 全部段落文本。"""
    from docx import Document  # type: ignore[import-untyped]

    document = Document(str(path))
    parts: list[str] = []
    for para in document.paragraphs:
        if para.text:
            parts.append(para.text)
    # 也提取表格中的文本
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    parts.append(cell.text)
    return "\n".join(parts)


def _read_text_safely(path: Path) -> str:
    """容错读取文本文件：utf-8 优先，失败回退 gbk。"""
    try:
        return path.read_text(encoding="utf-8")
    except (OSError, UnicodeDecodeError):
        try:
            return path.read_text(encoding="gbk")
        except (OSError, UnicodeDecodeError):
            return ""


def _build_suggestions(detected: list[str], clause_type: str | None) -> list[str]:
    """根据检测到的风险点生成建议。"""
    if not detected:
        return ["未检测到明显风险关键词，建议人工复核条款措辞。"]

    suggestions: list[str] = []
    suggestion_map: dict[str, str] = {
        "定金不退": "「定金不退」可能违反《民法典》定金规则，建议明确退还条件与情形。",
        "最终解释权": "「最终解释权」涉嫌排除消费者权利，违反《消费者权益保护法》第26条。",
        "单方变更": "「单方变更」可能加重对方风险，建议增加变更需双方协商一致条款。",
        "加重对方责任": "违约责任过重可能被法院调减，建议参照实际损失约定违约金。",
        "限制对方权利": "「视为放弃」等表述可能被认定为无效格式条款，建议双方权利对等。",
        "管辖不利": "管辖约定偏向一方可能在诉讼中被认定无效，建议约定被告所在地或合同履行地。",
        "自动续约": "「自动续约」应明确提前通知期与取消方式，保障对方知情权。",
        "没收财产": "「没收」「扣除全部」表述可能违法，建议改为按实际损失赔偿。",
        "格式条款排除": "「已充分理解」等格式条款排除表述在争议时举证力有限，建议补充实际说明义务。",
    }
    for risk in detected:
        if risk in suggestion_map:
            suggestions.append(suggestion_map[risk])
    if clause_type:
        suggestions.append(f"针对「{clause_type}」类条款，建议结合具体业务场景进一步审查。")
    return suggestions


__all__ = [
    "DocumentExtractResult",
    "ContractAnalysisResult",
    "extract_document",
    "analyze_contract_clause",
    "_RISK_CLAUSE_PATTERNS",
]
