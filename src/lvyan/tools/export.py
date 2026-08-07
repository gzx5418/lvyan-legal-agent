"""导出工具：将 Markdown 文本渲染为 DOCX（或降级为 .md）。

本模块是 SubTask 15.5 的实现，复用原 ``律言skill`` / ``law-consult-cn`` 中的
markdown_to_docx 转换逻辑，使用 python-docx 渲染标题、段落、列表、引用、
代码块。失败时优雅降级为 .md 文件。

支持基于官方模板（``template`` 参数指向 .docx 文件）的渲染：模板存在时
将模板复制后追加 Markdown 转换的段落。
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import Literal


from lvyan.tools.base import ToolResult

# ---------------------------------------------------------------------------
# 常量
# ---------------------------------------------------------------------------
ExportFormat = Literal["docx", "md"]

_HEADING_RE = re.compile(r"^(#{1,6})\s+(.+)$")
_ORDERED_RE = re.compile(r"^\s*\d+[.)]\s+(.+)$")
_UNORDERED_RE = re.compile(r"^\s*[-*+]\s+(.+)$")


# ---------------------------------------------------------------------------
# 返回模型
# ---------------------------------------------------------------------------
class ExportResult(ToolResult):
    """文书导出结果。"""

    output_path: str = ""
    format: ExportFormat = "docx"
    file_size: int = 0
    # success / error 由基类提供


# ---------------------------------------------------------------------------
# 公开工具
# ---------------------------------------------------------------------------
def render_docx(
    markdown_text: str,
    output_path: str,
    template: str | None = None,
) -> ExportResult:
    """将 Markdown 文本渲染为 DOCX 文档（失败时降级为 .md）。

    Args:
        markdown_text: Markdown 正文文本。
        output_path: 输出文件路径。若以 ``.docx`` 结尾则尝试 DOCX 渲染；
            若 python-docx 不可用或渲染失败，自动改写为同名 ``.md`` 降级输出。
        template: 可选的模板 .docx 路径。模板存在时复制模板再追加内容。

    Returns:
        ExportResult：含 output_path / format / file_size / success / error。
    """
    if not markdown_text:
        return ExportResult(
            tool_name="render_docx",
            success=False,
            error="markdown_text 不能为空",
            output_path=output_path or "",
        )
    if not output_path:
        return ExportResult(
            tool_name="render_docx",
            success=False,
            error="output_path 不能为空",
            output_path="",
        )

    out_path = Path(output_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)

    # 尝试 DOCX 渲染
    try:
        _render_docx_internal(markdown_text, out_path, template)
        size = out_path.stat().st_size if out_path.is_file() else 0
        return ExportResult(
            tool_name="render_docx",
            success=True,
            output_path=str(out_path),
            format="docx",
            file_size=size,
        )
    except Exception as exc:  # noqa: BLE001
        # 降级写 .md
        md_path = out_path.with_suffix(".md")
        try:
            md_path.write_text(markdown_text, encoding="utf-8")
            size = md_path.stat().st_size if md_path.is_file() else 0
            return ExportResult(
                tool_name="render_docx",
                success=True,
                output_path=str(md_path),
                format="md",
                file_size=size,
                error=f"DOCX 渲染失败，已降级为 .md：{exc}",
            )
        except Exception as exc2:  # noqa: BLE001
            return ExportResult(
                tool_name="render_docx",
                success=False,
                error=f"DOCX 渲染失败：{exc}；.md 降级亦失败：{exc2}",
                output_path=str(out_path),
                format="md",
                file_size=0,
            )


# ---------------------------------------------------------------------------
# 内部：DOCX 渲染
# ---------------------------------------------------------------------------
def _render_docx_internal(markdown_text: str, out_path: Path, template: str | None) -> None:
    """用 python-docx 将 Markdown 渲染为 DOCX。失败抛出异常。"""
    from docx import Document  # type: ignore[import-untyped]

    # 模板存在时基于模板，否则空白文档
    if template:
        tpl_path = Path(template)
        if tpl_path.is_file():
            document = Document(str(tpl_path))
        else:
            document = Document()
    else:
        document = Document()

    _setup_styles(document)

    lines = markdown_text.splitlines()
    in_code = False
    code_buffer: list[str] = []

    for line in lines:
        stripped = line.rstrip()

        if stripped.startswith("```"):
            if in_code:
                if code_buffer:
                    paragraph = document.add_paragraph()
                    run = paragraph.add_run("\n".join(code_buffer))
                    run.font.name = "Consolas"
                    code_buffer = []
                in_code = False
            else:
                in_code = True
            continue

        if in_code:
            code_buffer.append(stripped)
            continue

        if not stripped:
            continue

        heading = _HEADING_RE.match(stripped)
        if heading:
            level = min(len(heading.group(1)), 4)
            document.add_heading(_strip_inline_markdown(heading.group(2)), level=level)
            continue

        unordered = _UNORDERED_RE.match(stripped)
        if unordered:
            document.add_paragraph(
                _strip_inline_markdown(unordered.group(1)),
                style="List Bullet",
            )
            continue

        ordered = _ORDERED_RE.match(stripped)
        if ordered:
            document.add_paragraph(
                _strip_inline_markdown(ordered.group(1)),
                style="List Number",
            )
            continue

        if stripped.startswith(">"):
            try:
                document.add_paragraph(
                    _strip_inline_markdown(stripped.lstrip("> ")),
                    style="Intense Quote",
                )
            except KeyError:
                # 模板可能缺少 Intense Quote 样式，回退为普通段落
                document.add_paragraph(_strip_inline_markdown(stripped.lstrip("> ")))
            continue

        document.add_paragraph(_strip_inline_markdown(stripped))

    out_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(out_path))


def _setup_styles(document) -> None:
    """设置正文样式为 Arial / 宋体 12pt（如模板已有样式则保留）。"""
    try:
        from docx.shared import Pt  # type: ignore[import-untyped]
        from docx.oxml.ns import qn  # type: ignore[import-untyped]
    except ImportError:
        return

    try:
        normal = document.styles["Normal"]
        normal.font.name = "Arial"
        normal.font.size = Pt(12)
        rpr = normal.element.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = rpr.makeelement(qn("w:rFonts"), {})
            rpr.insert(0, rfonts)
        rfonts.set(qn("w:ascii"), "Arial")
        rfonts.set(qn("w:hAnsi"), "Arial")
        rfonts.set(qn("w:eastAsia"), "宋体")
    except Exception:
        # 样式设置失败不阻断渲染
        pass


def _strip_inline_markdown(text: str) -> str:
    """去除行内 Markdown 标记（`` ` `` / ``**`` / ``*`` / 链接）。"""
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"\*([^*]+)\*", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1 (\2)", text)
    return text.strip()


__all__ = [
    "ExportResult",
    "render_docx",
]
